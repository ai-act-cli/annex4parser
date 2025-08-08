# document_ingestion.py
"""Utilities for ingesting compliance documents into the database.

This module provides helper functions for extracting plain text
from PDF and DOCX files and storing them in the compliance
database.  It then maps the extracted content to relevant AI Act
rules using the keyword matcher from :mod:`mapper` and records
those mappings in the database.  The implementation is a
prototype; it does not handle scanned PDFs or complex formats,
but establishes the plumbing needed for further expansion.
"""

from pathlib import Path
from typing import Optional

import pdfplumber  # type: ignore
import docx  # type: ignore
from sqlalchemy.orm import Session

# Prefer the combined keyword/semantic matcher over the plain keyword
# matcher.  This import is kept on a separate line so that tools
# injecting automated fixes do not remove unused imports prematurely.
from .mapper.combined_mapper import combined_match_rules
from .models import Document, DocumentRuleMapping, Rule


def extract_text_from_pdf(pdf_path: Path) -> str:
    """Extract plain text from a PDF file (skips scanned images)."""
    text: list[str] = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text.append(page.extract_text() or "")
    return "\n".join(text)


def extract_text_from_docx(docx_path: Path) -> str:
    """Extract plain text from a DOCX file."""
    document = docx.Document(docx_path)
    return "\n".join(p.text for p in document.paragraphs)


def ingest_document(
    file_path: Path,
    db: Session,
    *,
    customer_id: Optional[str] = None,
    ai_system_name: Optional[str] = None,
    document_type: Optional[str] = None,
) -> Document:
    """Parse a document, store it in the database and map it to rules.

    Parameters
    ----------
    file_path:
        Path to the document to ingest.
    db:
        SQLAlchemy session.
    customer_id, ai_system_name, document_type:
        Optional metadata for the :class:`~models.Document` record.

    Returns
    -------
    Document
        The created database record.
    """
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        text = extract_text_from_pdf(file_path)
    elif suffix in {".docx", ".doc"}:
        text = extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    document = Document(
        customer_id=customer_id,
        filename=file_path.name,
        file_path=str(file_path),
        extracted_text=text,
        ai_system_name=ai_system_name,
        document_type=document_type,
    )

    db.add(document)
    db.flush()  # obtain ID for relationships

    # Use the combined matcher to map the document's text to relevant
    # rules.  Only section codes with a non‑zero combined score are
    # returned.  This provides a more nuanced confidence value by
    # blending keyword presence with semantic similarity.
    matches = combined_match_rules(db, text)
    for section_code, confidence in matches.items():
        rule = db.query(Rule).filter_by(section_code=section_code).first()
        if not rule:
            continue
        mapping = DocumentRuleMapping(
            document_id=document.id,
            rule_id=rule.id,
            confidence_score=confidence,
            mapped_by="auto",
        )
        db.add(mapping)

    db.commit()
    return document
